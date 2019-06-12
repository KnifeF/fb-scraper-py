import numpy as np
import pandas as pd
import glob
from plot_charts import *
from docx import Document
from docx.image.exceptions import UnrecognizedImageError
from program_constants import *
from docx.shared import Inches
from time import sleep
from PIL import Image

# -*- coding: UTF-8 -*-
__author__ = "KnifeF"


class FbProfile:

    def __init__(self, facebook_uid, user_name=None, profile_name=None):
        # build facebook profile object
        self.facebook_uid = facebook_uid
        self.user_name = user_name
        self.profile_name = profile_name

        self.about = None
        self.posts = []
        self.groups = []
        self.pages_liked = []
        self.friends = []
        self.followers = []
        self.stories_liked = []
        self.stories_commented = []
        self.photos_liked = []
        self.photos_commented = []
        self.videos_liked = []
        self.videos_commented = []
        self.apps_used = []
        self.common_groups = []
        self.common_pages_liked = []
        self.common_friends = []
        self.common_followers = []
        self.common_stories_liked = []
        self.common_stories_commented = []
        self.common_photos_liked = []
        self.common_photos_commented = []
        self.common_videos_liked = []
        self.common_videos_commented = []
        self.common_apps_used = []

        self.screenshots = []
        self.photos = []
        self.profile_photo = None
        self.pics_on_web = None
        self.found_on_web_screenshot = None
        self.interests_donut_path = None
        self.security_friends = None
        self.relevant_groups = None
        self.friends_donut_path = None
        self.groups_donut_path = None

        self.build_profile_from_files()

    def build_profile_from_files(self):
        # build profile obj parameters from csv files (using pandas)
        self.build_about_from_pandas()
        self.set_name_and_user()
        self.is_exist_on_facebook()
        self.set_profile_pics()
        if os.path.exists(SAVE_PATH):
            all_dirs = os.listdir(SAVE_PATH)
            for dir_name in all_dirs:
                dir_path = os.path.join(SAVE_PATH, dir_name)
                csv_files = glob.glob1(dir_path, "*.csv")
                for f_name in csv_files:
                    if self.facebook_uid in f_name and "#" in f_name:
                        split_f_name = f_name.split("#")
                        if len(split_f_name) > 1:
                            print(f_name)
                            df = pd.read_csv(os.path.join(dir_path, f_name))
                            if not df.empty:
                                data_lst = df.to_dict('records')
                                data_lst = nan_to_none(data_lst)

                                if split_f_name[0].isdigit():
                                    self.set_activity(split_f_name[1], data_lst)
                                    if (not self.profile_name) and ('Profile Name' in data_lst[0]):
                                        self.profile_name = data_lst[0]['Profile Name']
                                elif '_' in split_f_name[0]:
                                    uids = split_f_name[0].split('_')
                                    if len(uids) == 2 and (uids[0].isdigit() and uids[1].isdigit()):
                                        self.set_common_activity(split_f_name[1], data_lst)
            self.from_pages_to_donut()
            self.from_friends_to_donut()
            self.from_groups_to_donut()

    def build_about_from_pandas(self):
        # build a dictionary with details about profile from csv (using pandas)
        about_file = os.path.join(PATH_ABOUT, "about_targets.csv")
        # print(about_file)
        if os.path.exists(about_file):
            print(about_file)
            df = pd.read_csv(about_file)
            if (not df.empty) and ("UID" in df) and self.facebook_uid:
                count = 0
                for val in df["UID"]:
                    if self.facebook_uid in str(val):
                        self.about = df.to_dict('records')[count]
                        self.set_about_dict()
                        break
                    print(val)
                    count += 1
                """
                for index, row in df.iterrows():
                    if self.facebook_uid == str(row['UID']):
                        self.about = df.to_dict('records')[index]
                        self.set_about_dict()
                        break
                """
                # print(index)
                # loc_row = df.loc[df['UID'] == self.facebook_uid]

    def set_about_dict(self):
        # set 'about' keys and values (dict)
        if self.about:
            self.about.pop('Time Stamp', None)  # pop key from dict
            self.about['Exist On Facebook'] = None  # add new key to dict
            self.about['Stolen Photos'] = None  # add new key to dict
            self.about['Is Fake'] = None  # add new key to dict

            # pop keys from dict and append the details to a new key
            self.change_keys_to_key('Education', 6)
            self.change_keys_to_key('Work', 6)
            self.change_keys_to_key('Website', 3)
            self.change_keys_to_key('Contact Info', 3)
            self.change_keys_to_key('Basic Info', 3)
            self.change_keys_to_key('Family Member', 10)
            self.change_keys_to_key('Life Event', 10)
            self.change_keys_to_key('Other Name', 12)

            # change 'nan' values in dict to None
            for key, value in self.about.items():
                if (not value) or (value and str(value) == str(np.nan)):
                    self.about[key] = None

    def change_keys_to_key(self, key_name, count_limit):
        # pop keys from dict and append the details to a new key
        if key_name and count_limit and self.about and self.about['%s1' % key_name]:
            new_key = '%s List' % key_name
            self.about[new_key] = []
            if 'Other Name' not in key_name:
                for i in range(1, count_limit+1):
                    value = self.about.pop('%s%d' % (key_name, i), None)
                    if value and (str(value) != str(np.nan)):
                        self.about[new_key].append(value)
            else:
                other_names = [r"Nickname1", r"Nickname2", r"Maiden Name", r"Alternative Name",
                               r"Married Name", r"Father's Name", r"Birth Name", r"Former Name",
                               r"Name With Title", r"Other Name1", r"Other Name2", r"Other Name3"]
                for i in range(len(other_names)):
                    value = self.about.pop(other_names[i], None)
                    if value and (str(value) != str(np.nan)):
                        self.about[new_key].append(value)
            # print(self.about[new_key])

    def is_exist_on_facebook(self):
        # has the profile been found on facebook recently
        if self.about:
            exist_file = os.path.join(PATH_EXIST, "profile_exist_status.csv")
            if os.path.exists(exist_file):
                df = pd.read_csv(exist_file)
                if ("UID" in df) and ("Exists On Facebook" in df) and self.facebook_uid:
                    count = 0
                    for val in df["UID"]:
                        if self.facebook_uid in str(val):
                            self.about['Exists On Facebook'] = df["Exists On Facebook"][count]
                            break
                        count += 1
                    # loc_row = df.loc[df["UID"] == self.facebook_uid]
                    # if not loc_row.empty:
                    # self.about['Exists On Facebook'] = loc_row['Exists On Facebook'].values[0]

    def set_activity(self, query, data_lst):
        # set required list to obj parameter (related to profile's activity)
        if query == "groups":
            self.groups = data_lst
        elif query == "pages-liked":
            self.pages_liked = data_lst
        elif query == "friends":
            self.friends = data_lst
        elif query == "followers":
            self.followers = data_lst
        elif query == "stories-liked":
            self.stories_liked = data_lst
        elif query == "photos-liked":
            self.photos_liked = data_lst
        elif query == "videos-liked":
            self.videos_liked = data_lst
        elif query == "stories-commented":
            self.stories_commented = data_lst
        elif query == "photos-commented":
            self.photos_commented = data_lst
        elif query == "videos-commented":
            self.videos_commented = data_lst
        elif query == "apps-used":
            self.apps_used = data_lst
        elif query == "posts":
            self.posts = data_lst
            # self.set_hashtags()
            # self.set_likes_info('Likes')
            # self.set_likes_info('Comments')
            # self.set_likes_info('Shares')

    def set_common_activity(self, query, data_lst):
        # set required list to obj parameter (related to profile's common activity with others)
        if query == "groups":
            self.common_groups.append(data_lst)
        elif query == "pages-liked":
            self.common_pages_liked.append(data_lst)
        elif query == "friends":
            self.common_friends.append(data_lst)
        elif query == "followers":
            self.common_followers.append(data_lst)
        elif query == "stories-liked":
            self.common_stories_liked.append(data_lst)
        elif query == "photos-liked":
            self.common_photos_liked.append(data_lst)
        elif query == "videos-liked":
            self.common_videos_liked.append(data_lst)
        elif query == "stories-commented":
            self.common_stories_commented.append(data_lst)
        elif query == "photos-commented":
            self.common_photos_commented.append(data_lst)
        elif query == "videos-commented":
            self.common_videos_commented.append(data_lst)
        elif query == "apps-used":
            self.common_apps_used.append(data_lst)

    def set_profile_pics(self):
        # check if there are profile pics that have occurrences on the web,
        # and set paths for profile pic&screenshot
        pics_path = os.path.join(PATH_PHOTOS, self.facebook_uid)
        if os.path.exists(pics_path) and os.path.isdir(pics_path):
            prof_pics = glob.glob1(pics_path, "*.jpg")
            pics_on_web_files = glob.glob1(pics_path, "*.csv")

            if prof_pics:
                self.profile_photo = os.path.join(pics_path, prof_pics[0])
                print(self.profile_photo)
            if pics_on_web_files and ('profile_pics_on_web' in pics_on_web_files[0]):
                pics_on_web_path = os.path.join(pics_path, pics_on_web_files[0])
                df = pd.read_csv(pics_on_web_path)
                if (not df.empty) and 'Found On The Web?' in df:
                    count = 0
                    for row in list(df['Found On The Web?']):
                        if 'Not found on the web' in row:
                            count += 1
                        else:
                            if 'Found on the web with' in row:
                                self.pics_on_web = 'Found on the web (with Google)'
                            if 'Found on the web with TinEye' in row:
                                if self.pics_on_web:
                                    self.pics_on_web = self.pics_on_web.replace('(with Google)', '(with Google and TinEye)')
                                else:
                                    self.pics_on_web = 'Found on the web (with TinEye)'
                    if (not self.pics_on_web) and count > 0:
                        self.pics_on_web = "Not found on the web"
                found_on_web_screenshots = glob.glob1(pics_path, "*.png")
                if found_on_web_screenshots and ('pic_found' in found_on_web_screenshots[0]) \
                        and (self.facebook_uid in found_on_web_screenshots[0]):
                    self.found_on_web_screenshot = os.path.join(pics_path, found_on_web_screenshots[0])

    def set_name_and_user(self):
        # set profile name and username if required (when details are found in 'about' dict)
        if not (self.user_name and self.profile_name):
            if self.about:
                if ('User Name' in self.about) and ('Profile Name' in self.about) and ('Facebook' in self.about):
                    if (not self.profile_name) and self.about['Profile Name']:
                        self.profile_name = self.about['Profile Name']
                    if not self.user_name:
                        if self.about['User Name']:
                            self.user_name = self.about['User Name']
                        elif self.about['Facebook']:
                            self.user_name = self.about['Facebook'].replace('/', '')

    def str_of_list_by_q(self, x, q):
        # convert list with strings to one string, to add for docx report
        if x != -1:
            res = "%d) %s: " % (x, q)
        else:
            res = "%s: " % q
        if self.about and ('%s List' % q in self.about) and self.about['%s List' % q]:
            lst = self.about['%s List' % q]
            if lst:
                for i in range(len(lst)):
                    if i < len(ABC_STR):
                        res += "\n  %s) %s" % (ABC_STR[i], lst[i])
        else:
            res += 'X'
        return res

    def str_by_q(self, x, q):
        if x != -1:
            res = "%d) %s: " % (x, q)
        else:
            res = "%s: " % q
        if self.about and (q in self.about) and self.about[q]:
            res += self.about[q]
        else:
            res += 'X'
        return res

    def str_similarity_with_others(self, doc=None):
        # find common things between the profile with others count them
        res = ""
        params_lst = [self.common_groups, self.common_pages_liked, self.common_friends, self.common_followers,
                      self.common_photos_commented, self.common_stories_commented, self.common_videos_commented,
                      self.common_photos_liked, self.common_stories_liked, self.common_videos_liked,
                      self.common_apps_used]

        for i in range(len(params_lst)):
            describe_common = common_str_from_cond(i)
            if doc:
                res += "\n\nCommon %s: " % describe_common
            else:
                res += "\n%d) Common %s: " % (i+1, describe_common)
            count = 0
            if len(params_lst[i]) > 0:
                for j in range(len(params_lst[i])):
                    current_list = params_lst[i][j]
                    if current_list and isinstance(current_list, list) and len(current_list) > 0 and j < len(ABC_STR):
                        current_dict = current_list[0]
                        if current_dict and isinstance(current_dict, dict) and ('UID' in current_dict) \
                                and ('UID2' in current_dict):
                            uids = (str(current_dict['UID']), str(current_dict['UID2']))
                            if self.facebook_uid in uids:
                                profile_dsc = ""
                                if uids[0] != self.facebook_uid:
                                    if 'Profile Name' in current_dict:
                                        profile_dsc = current_dict['Profile Name2']+", "
                                    profile_dsc += '[UID - %s]' % uids[0]
                                else:
                                    if 'Profile Name2' in current_dict:
                                        profile_dsc = current_dict['Profile Name']+", "
                                    profile_dsc += '[UID - %s]' % uids[1]
                                    # print(profile_dsc)
                                if profile_dsc:
                                    res += "\n  %s) %d Common %s With %s" \
                                           % (ABC_STR[count], len(current_list), describe_common, profile_dsc)
                                    count += 1
            else:
                res += 'X'
        return res

    def from_pages_to_donut(self):
        # create a donut chart of page categories (save to .png file), from pages_liked
        if self.pages_liked:
            df = pd.DataFrame(self.pages_liked)
            if r"Page Category" in df:
                page_categories = df[r"Page Category"]
                if len(page_categories) > 0:
                    self.interests_donut_path = os.path.join(PATH_PAGES, self.facebook_uid+"#interests_donut.png")
                    relevant_categories, sum_categories = filter_page_categories(page_categories)
                    plot_donut_chart(relevant_categories, sum_categories, self.interests_donut_path, "Pages")

    def from_friends_to_donut(self):
        # create a donut chart of profile's friends (save to .png file)
        if self.friends:
            df = pd.DataFrame(self.friends)
            if r"Work" in df:
                friends_work = df[r"Work"]
                if len(friends_work) > 0:
                    self.friends_donut_path = os.path.join(PATH_FRIENDS, self.facebook_uid+"#friends_donut.png")
                    security_percent_str, relevant_categories, sum_categories = filter_by_work(friends_work)
                    self.security_friends = security_percent_str
                    plot_donut_chart(relevant_categories, sum_categories, self.friends_donut_path, "Friends")

    def from_groups_to_donut(self):
        # create a donut chart of profile's groups (save to .png file)
        if self.groups:
            df = pd.DataFrame(self.groups)
            if r"Group ID" in df:
                prof_groups = df[r"Group ID"]
                if len(prof_groups) > 0:
                    self.groups_donut_path = os.path.join(PATH_GROUPS, self.facebook_uid+"#groups_donut.png")
                    group_percent_str, relevant_groups, sum_groups = filter_groups(prof_groups)
                    if group_percent_str and relevant_groups and sum_groups:
                        self.relevant_groups = group_percent_str
                        plot_donut_chart(relevant_groups, sum_groups, self.groups_donut_path, "Groups")

    def __str__(self):
        res = "Facebook Profile Report\n***************\n\nFacebook Contact Info & Social Accounts:" + \
              "\n1) Facebook uid: "+str(self.facebook_uid).replace('None', 'X') + \
              "\n2) Facebook Username: "+str(self.user_name).replace('None', 'X') + \
              "\n3) FB Profile Name: "+str(self.profile_name).replace('None', 'X') + \
              "\n"+self.str_by_q(4, 'Email')+"\n"+self.str_by_q(5, 'Phone Number') + \
              "\n"+self.str_by_q(6, 'Instagram')+"\n"+self.str_by_q(7, 'Twitter') + \
              "\n"+self.str_by_q(8, 'Snapchat')+"\n"+self.str_by_q(9, 'Youtube') + \
              "\n"+self.str_by_q(10, 'LinkedIn')+"\n"+self.str_by_q(11, 'Skype') + \
              "\n"+self.str_by_q(12, 'Twitter')+"\n"+self.str_of_list_by_q(13, 'Website') + \
              "\n"+self.str_of_list_by_q(14, 'Contact Info')+"\n\nFacebook Details:" + \
              "\n"+self.str_by_q(1, 'Current City')+"\n"+self.str_by_q(2, 'Hometown') + \
              "\n"+self.str_of_list_by_q(3, 'Work')+"\n"+self.str_of_list_by_q(4, 'Education') + \
              "\n"+self.str_by_q(5, 'Birthday')+"\n"+self.str_by_q(6, 'Gender') + \
              "\n"+self.str_by_q(7, 'Interested In')+"\n"+self.str_by_q(8, 'Relationship Status') + \
              "\n"+self.str_by_q(9, 'Languages')+"\n"+self.str_by_q(10, 'Religious Views') + \
              "\n"+self.str_by_q(11, 'Political Views')+"\n"+self.str_of_list_by_q(12, 'Basic Info') + \
              "\n"+self.str_of_list_by_q(13, 'Other Name')+"\n"+self.str_of_list_by_q(14, 'Family Member') + \
              "\n"+self.str_of_list_by_q(15, 'Life Event')+"\n"+self.str_by_q(16, 'Intro Description') + \
              "\n"+self.str_of_list_by_q(17, 'Bio')+"\n"+self.str_by_q(18, 'Favourite Quotes') + \
              "\n"+self.str_by_q(19, 'Exists On Facebook')+"\n"+self.str_by_q(20, 'Has Phone?') + \
              "\n\nFacebook Activity:"+"\n1) Groups: "+str(len(self.groups)) + \
              "\n2) Pages-liked: "+str(len(self.pages_liked))+"\n3) Friends: "+str(len(self.friends)) + \
              "\n4) Followers: "+str(len(self.followers))+"\n5) photos-commented: "+str(len(self.photos_commented)) + \
              "\n6) stories-commented: "+str(len(self.stories_commented)) + \
              "\n7) videos-commented: "+str(len(self.videos_commented))+"\n8) photos-liked : "+str(len(self.photos_liked)) + \
              "\n9) stories-liked: "+str(len(self.stories_liked))+"\n10) videos-liked : "+str(len(self.videos_liked)) + \
              "\n11) apps-used: "+str(len(self.apps_used)) + \
              "\n\nSimilarity With Other Profiles:"+"\n"+self.str_similarity_with_others()
        return res

    def set_data_to_report(self):
        # return lists with profile's data for docx report
        contact_info = ["Facebook uid: "+str(self.facebook_uid).replace('None', 'X'),
                        "Facebook Username: "+str(self.user_name).replace('None', 'X'),
                        "FB Profile Name: "+str(self.profile_name).replace('None', 'X'),
                        self.str_by_q(-1, 'Email'), self.str_by_q(-1, 'Phone Number'),
                        self.str_by_q(-1, 'Instagram'), self.str_by_q(-1, 'Twitter'),
                        self.str_by_q(-1, 'Snapchat'), self.str_by_q(-1, 'Youtube'),
                        self.str_by_q(-1, 'LinkedIn'), self.str_by_q(-1, 'Skype'),
                        self.str_of_list_by_q(-1, 'Website'), self.str_of_list_by_q(-1, 'Contact Info')]
        facebook_details = [self.str_by_q(-1, 'Current City'), self.str_by_q(-1, 'Hometown'),
                            self.str_of_list_by_q(-1, 'Work'), self.str_of_list_by_q(-1, 'Education'),
                            self.str_by_q(-1, 'Birthday'), self.str_by_q(-1, 'Gender'),
                            self.str_by_q(-1, 'Interested In'), self.str_by_q(-1, 'Relationship Status'),
                            self.str_by_q(-1, 'Languages'), self.str_by_q(-1, 'Religious Views'),
                            self.str_by_q(-1, 'Political Views'), self.str_of_list_by_q(-1, 'Basic Info'),
                            self.str_of_list_by_q(-1, 'Other Name'), self.str_of_list_by_q(-1, 'Family Member'),
                            self.str_of_list_by_q(-1, 'Life Event'), self.str_by_q(-1, 'Intro Description'),
                            self.str_by_q(-1, 'Bio'), self.str_by_q(-1, 'Favourite Quotes'),
                            self.str_by_q(-1, 'Exists On Facebook'), self.str_by_q(-1, 'Has Phone?')]
        facebook_activity = ["Groups: "+str(len(self.groups)), "Pages: "+str(len(self.pages_liked)),
                             "Friends: "+str(len(self.friends)), "Followers: "+str(len(self.followers)),
                             "photos-commented: "+str(len(self.photos_commented)),
                             "stories-commented: "+str(len(self.stories_commented)),
                             "videos-commented: "+str(len(self.videos_commented)),
                             "photos-liked : "+str(len(self.photos_liked)), "stories-liked: "+str(len(self.stories_liked)),
                             "videos-liked : "+str(len(self.videos_liked)), "apps-used: "+str(len(self.apps_used))]
        return contact_info, facebook_details, facebook_activity

    def profile_to_docx_report(self):
        # create docx report from profile's relevant parameters
        create_dir(PATH_REPORTS)
        if self.profile_not_empty():
            heading = self.facebook_uid
            if self.profile_name:
                heading = self.profile_name

            contact_info, facebook_details, facebook_activity = self.set_data_to_report()
            document = Document()
            document.add_heading("FB Profile Report - %s" % heading, 1)
            if self.profile_photo and '.jpg' in self.profile_photo:
                # document.add_picture(self.profile_photo, width=Inches(1.25))
                pic_added = try_add_pic(document, self.profile_photo, width=Inches(2.25))
                if not pic_added:
                    try_add_pic(document, self.profile_photo, width=Inches(2.25), re_save=True)
            document.add_heading("Facebook Contact Info & Social Accounts:", 2)
            add_paragraph_from_list(document, contact_info)
            document.add_heading("Facebook Details:", 2)
            add_paragraph_from_list(document, facebook_details)
            add_paragraph_from_list(document, [('Security Friends: %s' % self.security_friends).replace('None', "X")])
            add_paragraph_from_list(document, [('Relevant Groups: %s' % self.relevant_groups).replace('None', "X")])
            add_paragraph_from_list(document, [('Photo Status: %s' % self.pics_on_web).replace('None', "X")])
            document.add_heading("Facebook Activity:", 2)
            add_paragraph_from_list(document, facebook_activity)
            document.add_heading("Similarity With Other Profiles:", 2)
            similarity_with_others = self.str_similarity_with_others(doc=document)
            add_paragraph_from_list(document, similarity_with_others.split('\n\n'))
            count = 0
            if self.found_on_web_screenshot:
                count += 1
                document.add_page_break()
                document.add_heading("Appendix %d - the photo is found on web:" % count, 2)
                try_add_pic(document, self.found_on_web_screenshot, width=Inches(6.25))
                # document.add_picture(self.found_on_web_screenshot, width=Inches(6.25))
            if self.groups_donut_path and os.path.exists(self.interests_donut_path):
                count += 1
                document.add_heading("Appendix %d - profile's groups:" % count, 2)
                try_add_pic(document, self.groups_donut_path, width=Inches(6.25))
            if self.interests_donut_path and os.path.exists(self.interests_donut_path):
                count += 1
                document.add_heading("Appendix %d - profile's interests (pages categories):" % count, 2)
                try_add_pic(document, self.interests_donut_path, width=Inches(6.25))
            if self.friends_donut_path and os.path.exists(self.friends_donut_path):
                count += 1
                document.add_heading("Appendix %d - profile's friends:" % count, 2)
                try_add_pic(document, self.friends_donut_path, width=Inches(6.25))
            document.save(os.path.join(PATH_REPORTS, '%s report.docx' % heading))

    def profile_not_empty(self):
        # check if profile obj has enough data required to create a report
        if not self.facebook_uid:
            return False
        return (self.user_name or self.profile_name or self.about or self.posts or self.groups
                or self.pages_liked or self.friends or self.followers or self.stories_liked
                or self.stories_commented or self.photos_liked or self.photos_commented
                or self.videos_liked or self.videos_commented or self.apps_used
                or self.common_groups or self.common_pages_liked or self.common_friends
                or self.common_followers or self.common_stories_liked or self.common_stories_commented
                or self.common_photos_liked or self.common_photos_commented or self.common_videos_liked
                or self.common_videos_commented or self.common_apps_used or self.screenshots
                or self.photos or self.profile_photo or self.pics_on_web or self.found_on_web_screenshot)


def try_add_pic(document, path, width=None, re_save=False):
    # add picture to a .docx file, by the given path (handling exceptions)
    # when 're_save' is True - try to fix a problem with adding picture to docx (with PIL)
    try:
        if re_save:
            img = Image.open(path)  # open image by path
            img.save(path, "JPEG")  # save opened image on the same name (replace the old file)
            sleep(1)
        if width:
            document.add_picture(path, width=width)  # add picture to docx with given width
        else:
            document.add_picture(path)  # add picture to docx
        return True
    except UnrecognizedImageError:  # 'UnrecognizedImageError' (couldn't add the picture to docx)
        print("UnrecognizedImageError")
        pass
    except IOError:  # 'IOError' - If the file cannot be found or opened
        print("IOError")
        pass
    return False


def add_paragraph_from_list(document, data_lst):
    for i in data_lst:
        if i:
            paragraph = document.add_paragraph(style='List Bullet')
            if ':' in i:
                paragraph.add_run(i.split(': ')[0]+": ").bold = True
                paragraph.add_run(i.split(': ')[1])


def common_str_from_cond(x):
    # return string by number
    return {
        0: "Groups", 1: "Pages", 2: "Friends", 3: "Followers",
        4: "photos-commented", 5: "stories-commented", 6: "videos-commented",
        7: "photos-liked", 8: "stories-liked", 9: "videos-liked",
        10: "apps-used"
    }[x]


def nan_to_none(data_lst):
    # change all np.nan values to 'None'
    if data_lst:
        for index in data_lst:
            # change 'nan' values in dict to None
            for key, value in index.items():
                if (not value) or (value and str(value) == str(np.nan)):
                    index[key] = None
    return data_lst
